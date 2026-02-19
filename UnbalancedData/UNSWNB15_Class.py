
import os, time, warnings, sys, io, traceback
from datetime import datetime
from pathlib import Path

import numpy as np
import pandas as pd

from sklearn.preprocessing import StandardScaler, LabelEncoder, label_binarize
from sklearn.metrics import (
    classification_report, accuracy_score, balanced_accuracy_score,
    precision_recall_fscore_support, confusion_matrix,
    roc_auc_score, average_precision_score,
    matthews_corrcoef, cohen_kappa_score
)
from sklearn.feature_selection import VarianceThreshold
from sklearn.svm import LinearSVC
from sklearn.calibration import CalibratedClassifierCV
from sklearn.neural_network import MLPClassifier
from sklearn.exceptions import ConvergenceWarning


# =========================================================
# (OPTIONNEL) — WORD EXPORT
# =========================================================
DOCX_OK = True
try:
    from docx import Document
    from docx.shared import Pt
except Exception:
    DOCX_OK = False


# =========================================================
# CAPTURE CONSOLE (TOUT CE QUI EST PRINT)
# =========================================================
class _TeeStdout:
    def __init__(self, original_stdout, buffer):
        self.original_stdout = original_stdout
        self.buffer = buffer

    def write(self, s):
        self.original_stdout.write(s)
        self.buffer.write(s)

    def flush(self):
        self.original_stdout.flush()

_console_buffer = io.StringIO()
_original_stdout = sys.stdout
sys.stdout = _TeeStdout(_original_stdout, _console_buffer)


# =========================================================
# STOCKAGE STRUCTURÉ DES RÉSULTATS 
# =========================================================
RESULTS = {"multiclass": [], "binary": []}

def _store_result(task, model, metrics, extra=None):
    row = {"task": task, "model": model}
    row.update(metrics or {})
    if extra:
        row.update(extra)
    if task == "binary":
        RESULTS["binary"].append(row)
    else:
        RESULTS["multiclass"].append(row)


# =========================================================
# SUPPRESSION DES WARNINGS (MLP / SVM)
# =========================================================
warnings.filterwarnings("ignore", category=ConvergenceWarning)
warnings.filterwarnings("ignore", message=".*Stochastic Optimizer.*")
warnings.filterwarnings("ignore", message=".*Liblinear failed to converge.*")


# =========================================================
# PARAMÈTRES (UNSW-NB15)
# =========================================================
BASE_DIR     = r"D:\MEMOIRE Merveilles\DataSet\UNSW_NB15"
PARQUET_NAME = "UNSW_NB15.parquet"

LABEL_COL_CANDIDATES = ["attack_cat", "Attack_cat", "Label", "label", "class", "Class"]
BENIGN_LABEL_DEFAULT = "Normal"

N_SAMPLE        = 5000
TOP_K_ATTACKS   = 10
BENIGN_KEEP     = 800
MIN_PER_CLASS   = 50

TRAIN_RATIO     = 0.7

VAR_THRESHOLD   = 0.01
RANDOM_STATE    = 42
np.random.seed(RANDOM_STATE)

ELM_HIDDEN = 64
NN_BIN = dict(hidden_layer_sizes=(64, 32), activation="tanh", max_iter=200, random_state=RANDOM_STATE)
NN_MUL = dict(hidden_layer_sizes=(128, 64), activation="tanh", max_iter=250, random_state=RANDOM_STATE)

SVM_MAX_ITER = 5000


# =========================================================
# OUTILS ROBUSTES
# =========================================================
def safe_len(x):
    try:
        return len(x)
    except Exception:
        return 1

def ensure_list(x):
    if isinstance(x, list):
        return x
    if isinstance(x, (pd.Index, pd.Series, np.ndarray, tuple)):
        return list(x)
    return [x]

def ensure_series(x):
    if isinstance(x, pd.Series):
        return x
    if isinstance(x, dict):
        return pd.Series(x)
    if isinstance(x, (pd.Index, list, tuple, np.ndarray)):
        return pd.Series(list(x))
    return pd.Series([x])

def pick_label_col(df):
    for c in LABEL_COL_CANDIDATES:
        if c in df.columns:
            return c
    raise ValueError(f"[ERREUR] Aucune colonne label trouvée parmi: {LABEL_COL_CANDIDATES}")

def detect_time_col(df):
    candidates = ["Timestamp", "timestamp", "stime", "Stime", "StartTime", "start_time", "date", "Date", "time", "Time"]
    for c in candidates:
        if c in df.columns:
            try:
                tmp = pd.to_datetime(df[c], errors="coerce")
                if tmp.isna().mean() < 0.2:
                    df[c] = tmp
                    return c
            except Exception:
                return c
    return None


# =========================================================
# AFFICHAGE CONSOLE (épuré)
# =========================================================
def h1(title):
    print("\n" + "="*100)
    print(title)
    print("="*100)

def h2(title):
    print("\n" + "-"*100)
    print(title)
    print("-"*100)

def timed_fit(fn):
    t0 = time.perf_counter()
    fn()
    return time.perf_counter() - t0

def show_counts(s, title, max_rows=30):
    print(f"\n[{title}]")
    vc = pd.Series(s).value_counts()
    if safe_len(vc) > max_rows:
        print(vc.head(max_rows).to_string())
        print(f"... ({safe_len(vc)-max_rows} autres classes)")
    else:
        print(vc.to_string())

def softmax(z):
    z = z - np.max(z, axis=1, keepdims=True)
    e = np.exp(z)
    return e / (np.sum(e, axis=1, keepdims=True) + 1e-12)


# =========================================================
# MÉTRIQUES
# =========================================================
def compute_summary_metrics(y_true, y_pred):
    acc  = accuracy_score(y_true, y_pred)
    bacc = balanced_accuracy_score(y_true, y_pred)
    mcc  = matthews_corrcoef(y_true, y_pred)
    kappa= cohen_kappa_score(y_true, y_pred)

    prec_m, rec_m, f1_m, _ = precision_recall_fscore_support(y_true, y_pred, average="macro", zero_division=0)
    prec_w, rec_w, f1_w, _ = precision_recall_fscore_support(y_true, y_pred, average="weighted", zero_division=0)

    return dict(
        Acc=float(acc), BAcc=float(bacc), MCC=float(mcc), Kappa=float(kappa),
        Macro_P=float(prec_m), Macro_R=float(rec_m), Macro_F1=float(f1_m),
        Weighted_P=float(prec_w), Weighted_R=float(rec_w), Weighted_F1=float(f1_w),
    )

def print_summary(y_true, y_pred, title=""):
    m = compute_summary_metrics(y_true, y_pred)
    print(f"[{title}] Acc={m['Acc']:.4f} | BAcc={m['BAcc']:.4f} | MCC={m['MCC']:.4f} | Kappa={m['Kappa']:.4f}")
    print(f"[{title}] Macro P/R/F1={m['Macro_P']:.4f}/{m['Macro_R']:.4f}/{m['Macro_F1']:.4f} | "
          f"Weighted P/R/F1={m['Weighted_P']:.4f}/{m['Weighted_R']:.4f}/{m['Weighted_F1']:.4f}")
    return m

def print_binary_rates(y_true, y_pred, title=""):
    cm = confusion_matrix(y_true, y_pred, labels=[0,1])
    if cm.size != 4:
        print(f"[{title}] Confusion Matrix:\n{cm}")
        return cm, None
    tn, fp, fn, tp = cm.ravel()
    tpr = tp/(tp+fn+1e-12); fpr = fp/(fp+tn+1e-12)
    tnr = tn/(tn+fp+1e-12); fnr = fn/(fn+tp+1e-12)
    print(f"[{title}] Confusion Matrix:\n{cm}")
    print(f"[{title}] TPR={tpr:.4f} | FPR={fpr:.4f} | TNR={tnr:.4f} | FNR={fnr:.4f}")
    return cm, dict(TPR=float(tpr), FPR=float(fpr), TNR=float(tnr), FNR=float(fnr))

def auc_binary(y_true, y_score, title=""):
    if safe_len(np.unique(y_true)) != 2:
        return None
    try:
        roc = roc_auc_score(y_true, y_score)
        ap  = average_precision_score(y_true, y_score)
        print(f"[{title}] ROC-AUC={roc:.4f} | PR-AUC(AP)={ap:.4f}")
        return dict(ROC_AUC=float(roc), PR_AUC=float(ap))
    except Exception:
        return None

def auc_multiclass(y_true, y_proba, n_classes, title=""):
    out = {}
    try:
        roc_macro = roc_auc_score(y_true, y_proba, multi_class="ovr", average="macro")
        roc_w     = roc_auc_score(y_true, y_proba, multi_class="ovr", average="weighted")
        print(f"[{title}] ROC-AUC OVR macro={roc_macro:.4f} | weighted={roc_w:.4f}")
        out["ROC_AUC_macro"] = float(roc_macro)
        out["ROC_AUC_weighted"] = float(roc_w)
    except Exception:
        pass

    try:
        Y = label_binarize(y_true, classes=np.arange(n_classes))
        ap_macro = average_precision_score(Y, y_proba, average="macro")
        ap_w     = average_precision_score(Y, y_proba, average="weighted")
        print(f"[{title}] PR-AUC macro={ap_macro:.4f} | weighted={ap_w:.4f}")
        out["PR_AUC_macro"] = float(ap_macro)
        out["PR_AUC_weighted"] = float(ap_w)
    except Exception:
        pass

    return out if out else None


# =========================================================
# ELM
# =========================================================
class ELM:
    def __init__(self, n_hidden=64, random_state=42):
        self.rng = np.random.default_rng(random_state)
        self.n_hidden = n_hidden

    def fit(self, X, Y_onehot):
        X = X.astype(np.float32)
        self.W = self.rng.standard_normal((X.shape[1], self.n_hidden)).astype(np.float32)
        self.b = self.rng.standard_normal(self.n_hidden).astype(np.float32)
        H = np.tanh(X @ self.W + self.b)
        self.beta = np.linalg.pinv(H) @ Y_onehot
        return self

    def decision_function(self, X):
        X = X.astype(np.float32)
        H = np.tanh(X @ self.W + self.b)
        return H @ self.beta

    def predict(self, X):
        return np.argmax(self.decision_function(X), axis=1)

    def predict_proba(self, X):
        return softmax(self.decision_function(X))


# =========================================================
# Échantillonnage étalé par classe
# =========================================================
def sample_spread_in_class(sub_df, n_take, time_col=None):
    if time_col is not None:
        sub_df = sub_df.sort_values(time_col)
    n = len(sub_df)
    if n_take >= n:
        return sub_df.copy()
    idx = np.linspace(0, n-1, n_take, dtype=int)
    return sub_df.iloc[idx].copy()


# =========================================================
# ÉCHANTILLON "ÉTALÉ + MAJORITAIRE"
# =========================================================
def build_majority_spread_sample(df, label_col, benign_label,
                                 n_total=5000, top_k_attacks=8,
                                 benign_keep=200, min_per_class=50,
                                 time_col=None):

    vc = pd.Series(df[label_col]).value_counts()

    attacks = vc.drop(index=[benign_label], errors="ignore")
    attacks = ensure_series(attacks)

    top_attacks = ensure_list(attacks.head(top_k_attacks).index)
    classes_keep = [benign_label] + top_attacks

    vc_keep = vc[vc.index.isin(classes_keep)]
    vc_keep = ensure_series(vc_keep)

    if benign_label not in vc_keep.index:
        vc_keep.loc[benign_label] = 0

    n_benign = int(min(benign_keep, n_total))
    if int(vc_keep.loc[benign_label]) == 0:
        n_benign = 0

    remaining = int(n_total - n_benign)

    attack_supports = vc_keep.drop(index=[benign_label], errors="ignore")
    attack_supports = ensure_series(attack_supports)
    total_attack_support = int(attack_supports.sum()) if safe_len(attack_supports) else 0

    parts = []
    if n_benign > 0:
        sub = df[df[label_col] == benign_label]
        parts.append(sample_spread_in_class(sub, min(n_benign, len(sub)), time_col=time_col))

    if remaining <= 0 or total_attack_support <= 0:
        out = pd.concat(parts, axis=0).reset_index(drop=True) if parts else df.iloc[0:0].copy()
        return out, top_attacks

    alloc = {}
    for cls, supp in attack_supports.items():
        supp_i = int(supp)
        alloc[cls] = min(int(min_per_class), supp_i)

    base = int(sum(alloc.values()))

    if base > remaining:
        scale = remaining / max(base, 1)
        for cls in list(alloc.keys()):
            alloc[cls] = max(1, int(alloc[cls] * scale))
    else:
        left = remaining - base
        if left > 0:
            weights = attack_supports / max(total_attack_support, 1)
            extra = (weights * left).astype(int)
            extra = ensure_series(extra)

            diff = int(left - int(extra.sum()))
            order = ensure_list(attack_supports.sort_values(ascending=False).index)

            if diff > 0 and safe_len(order) > 0:
                for i in range(diff):
                    extra.loc[order[i % safe_len(order)]] += 1

            for cls in alloc:
                cls_supp = int(attack_supports.get(cls, 0))
                alloc[cls] = min(cls_supp, int(alloc[cls] + int(extra.get(cls, 0))))

    for cls, n_take in alloc.items():
        sub = df[df[label_col] == cls]
        if len(sub) == 0:
            continue
        parts.append(sample_spread_in_class(sub, min(int(n_take), len(sub)), time_col=time_col))

    out = pd.concat(parts, axis=0).reset_index(drop=True)
    if time_col is not None and (not out.empty):
        out = out.sort_values(time_col).reset_index(drop=True)

    return out, top_attacks


# =========================================================
# Split temporel par classe (70/30) puis concat
# =========================================================
def classwise_time_split(df, label_col, time_col=None, train_ratio=0.7, min_train=1, min_test=1):
    train_parts, test_parts = [], []
    for cls in ensure_list(pd.Series(df[label_col]).unique()):
        sub = df[df[label_col] == cls].copy()
        if time_col is not None:
            sub = sub.sort_values(time_col)
        n = len(sub)
        if n == 0:
            continue
        split = int(train_ratio * n)
        if n >= (min_train + min_test):
            split = max(min_train, min(split, n - min_test))
            train_parts.append(sub.iloc[:split])
            test_parts.append(sub.iloc[split:])
        else:
            train_parts.append(sub)

    df_train = pd.concat(train_parts, axis=0).reset_index(drop=True) if train_parts else df.iloc[0:0].copy()
    df_test  = pd.concat(test_parts, axis=0).reset_index(drop=True) if test_parts else df.iloc[0:0].copy()

    if time_col is not None and (not df_train.empty):
        df_train = df_train.sort_values(time_col).reset_index(drop=True)
    if time_col is not None and (not df_test.empty):
        df_test  = df_test.sort_values(time_col).reset_index(drop=True)
    return df_train, df_test


# =========================================================
# Build X (fit sur train seulement)
# =========================================================
def build_X(df_train, df_test, time_col=None):
    X_train_df = df_train.select_dtypes(include=[np.number]).copy()
    X_test_df  = df_test.select_dtypes(include=[np.number]).copy()

    if time_col is not None:
        X_train_df = X_train_df.drop(columns=[time_col], errors="ignore")
        X_test_df  = X_test_df.drop(columns=[time_col], errors="ignore")

    common = X_train_df.columns.intersection(X_test_df.columns)
    X_train_df = X_train_df[common]
    X_test_df  = X_test_df[common]

    scaler = StandardScaler()
    X_train = scaler.fit_transform(X_train_df)
    X_test  = scaler.transform(X_test_df)

    selector = VarianceThreshold(threshold=VAR_THRESHOLD)
    X_train = selector.fit_transform(X_train)
    X_test  = selector.transform(X_test)
    return X_train, X_test


# =========================================================
# RUN MODELS
# =========================================================
def run_binary_models(X_train, X_test, y_train, y_test):
    def report(yp):
        return classification_report(
            y_test, yp,
            labels=[0,1],
            target_names=["BENIGN","ATTACK"],
            zero_division=0
        )

    # ELM
    h2("ELM — BINAIRE")
    elm = ELM(n_hidden=ELM_HIDDEN, random_state=RANDOM_STATE)
    t = timed_fit(lambda: elm.fit(X_train, np.eye(2)[y_train]))
    proba = elm.predict_proba(X_test)[:, 1]
    pred  = (proba >= 0.5).astype(int)
    print(f"Temps={t:.2f}s")
    met = print_summary(y_test, pred, "ELM BIN")
    _, rates = print_binary_rates(y_test, pred, "ELM BIN")
    aucs = auc_binary(y_test, proba, "ELM BIN")
    print(report(pred))
    _store_result("binary", "ELM", {**met, **(rates or {})}, extra={"TrainTime_s": float(t), **(aucs or {})})

    # NN
    h2("NN — BINAIRE")
    nn = MLPClassifier(**NN_BIN)
    t = timed_fit(lambda: nn.fit(X_train, y_train))
    proba = nn.predict_proba(X_test)[:, 1]
    pred  = (proba >= 0.5).astype(int)
    print(f"Temps={t:.2f}s")
    met = print_summary(y_test, pred, "NN BIN")
    _, rates = print_binary_rates(y_test, pred, "NN BIN")
    aucs = auc_binary(y_test, proba, "NN BIN")
    print(report(pred))
    _store_result("binary", "NN", {**met, **(rates or {})}, extra={"TrainTime_s": float(t), **(aucs or {})})

    # SVM
    h2("SVM — BINAIRE")
    svm = CalibratedClassifierCV(
        LinearSVC(random_state=RANDOM_STATE, max_iter=SVM_MAX_ITER),
        method="sigmoid"
    )
    t = timed_fit(lambda: svm.fit(X_train, y_train))
    proba = svm.predict_proba(X_test)[:, 1]
    pred  = (proba >= 0.5).astype(int)
    print(f"Temps={t:.2f}s")
    met = print_summary(y_test, pred, "SVM BIN")
    _, rates = print_binary_rates(y_test, pred, "SVM BIN")
    aucs = auc_binary(y_test, proba, "SVM BIN")
    print(report(pred))
    _store_result("binary", "SVM", {**met, **(rates or {})}, extra={"TrainTime_s": float(t), **(aucs or {})})


def run_multiclass_models(X_train, X_test, y_train, y_test, class_names):
    n_classes = safe_len(class_names)
    labels_all = np.arange(n_classes)

    def report(yp):
        return classification_report(
            y_test, yp,
            labels=labels_all,
            target_names=class_names,
            zero_division=0
        )

    # ELM
    h2("ELM — MULTICLASSE")
    elm = ELM(n_hidden=ELM_HIDDEN, random_state=RANDOM_STATE)
    t = timed_fit(lambda: elm.fit(X_train, np.eye(n_classes)[y_train]))
    proba = elm.predict_proba(X_test)
    pred  = np.argmax(proba, axis=1)
    print(f"Temps={t:.2f}s")
    met = print_summary(y_test, pred, "ELM MUL")
    aucs = auc_multiclass(y_test, proba, n_classes, "ELM MUL")
    print(report(pred))
    _store_result("multiclass", "ELM", met, extra={"TrainTime_s": float(t), **(aucs or {})})

    # NN
    h2("NN — MULTICLASSE")
    nn = MLPClassifier(**NN_MUL)
    t = timed_fit(lambda: nn.fit(X_train, y_train))
    proba = nn.predict_proba(X_test)
    pred  = np.argmax(proba, axis=1)
    print(f"Temps={t:.2f}s")
    met = print_summary(y_test, pred, "NN MUL")
    aucs = auc_multiclass(y_test, proba, n_classes, "NN MUL")
    print(report(pred))
    _store_result("multiclass", "NN", met, extra={"TrainTime_s": float(t), **(aucs or {})})

    # SVM
    h2("SVM — MULTICLASSE")
    svm = CalibratedClassifierCV(
        LinearSVC(random_state=RANDOM_STATE, max_iter=SVM_MAX_ITER),
        method="sigmoid"
    )
    t = timed_fit(lambda: svm.fit(X_train, y_train))
    proba = svm.predict_proba(X_test)
    pred  = np.argmax(proba, axis=1)
    print(f"Temps={t:.2f}s")
    met = print_summary(y_test, pred, "SVM MUL")
    aucs = auc_multiclass(y_test, proba, n_classes, "SVM MUL")
    print(report(pred))
    _store_result("multiclass", "SVM", met, extra={"TrainTime_s": float(t), **(aucs or {})})


# =========================================================
# RAPPORT (BUREAU)
# =========================================================
def get_desktop_dir():
    home = Path.home()
    d1 = home / "Desktop"
    if d1.exists():
        return d1
    d2 = home / "Bureau"
    if d2.exists():
        return d2
    return home

def _rank_and_interpret(rows, task_name):
    if not rows:
        return pd.DataFrame(), f"Aucun résultat enregistré pour la tâche {task_name}."

    df = pd.DataFrame(rows).copy()

    sort_cols = []
    if "Macro_F1" in df.columns: sort_cols.append("Macro_F1")
    if "BAcc" in df.columns: sort_cols.append("BAcc")
    if "Acc" in df.columns: sort_cols.append("Acc")

    df_sorted = df.sort_values(sort_cols, ascending=False).reset_index(drop=True) if sort_cols else df
    best = df_sorted.iloc[0].to_dict()

    text = []
    text.append(f"Interprétation (tâche {task_name}) :")
    text.append(
        "Pour les IDS déséquilibrés, la Macro-F1 et la Balanced Accuracy sont plus informatives que l’Accuracy seule."
    )
    text.append(
        f"Meilleur modèle observé : {best.get('model','?')} "
        f"(Macro_F1={best.get('Macro_F1', float('nan')):.4f}, "
        f"BAcc={best.get('BAcc', float('nan')):.4f}, "
        f"Acc={best.get('Acc', float('nan')):.4f})."
    )

    if task_name == "binaire":
        if "TPR" in best and "FPR" in best:
            text.append(f"TPR={best['TPR']:.4f} (détection attaques), FPR={best['FPR']:.4f} (fausses alertes).")
        if "ROC_AUC" in best:
            text.append(f"ROC-AUC={best['ROC_AUC']:.4f}.")
        if "PR_AUC" in best:
            text.append(f"PR-AUC(AP)={best['PR_AUC']:.4f}.")
    else:
        if "ROC_AUC_macro" in best:
            text.append(f"ROC-AUC OVR macro={best['ROC_AUC_macro']:.4f}.")
        if "PR_AUC_macro" in best:
            text.append(f"PR-AUC macro={best['PR_AUC_macro']:.4f}.")

    return df_sorted, "\n".join(text)

def save_report(console_text, results_dict, meta):
    desktop = get_desktop_dir()
    out_dir = desktop / "RESULTATS_UNSWNB15_CLASSIQUE_WORD"
    out_dir.mkdir(parents=True, exist_ok=True)

    ts = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    base_name = out_dir / f"UNSWNB15_CLASSIQUE_ELM_NN_SVM_{ts}"

    df_mc, txt_mc = _rank_and_interpret(results_dict.get("multiclass", []), "multiclasse")
    df_bi, txt_bi = _rank_and_interpret(results_dict.get("binary", []), "binaire")

    if DOCX_OK:
        doc = Document()
        doc.add_heading("UNSW-NB15 — Rapport de résultats (ELM / NN / SVM)", level=1)
        doc.add_paragraph(f"Généré le : {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Dataset : {meta.get('dataset_path','')}")
        doc.add_paragraph(f"Parquet : {meta.get('parquet_name','')}")
        doc.add_paragraph(f"Label col : {meta.get('label_col','')} | Benign : {meta.get('benign_label','')}")
        doc.add_paragraph(f"N_SAMPLE={meta.get('n_sample','')} | TRAIN_RATIO={meta.get('train_ratio','')} | VAR_THRESHOLD={meta.get('var_threshold','')}")
        if meta.get("time_col", None) is not None:
            doc.add_paragraph(f"Time col utilisée : {meta.get('time_col')}")
        else:
            doc.add_paragraph("Time col : non détectée (ordre des lignes utilisé).")

        doc.add_paragraph("")

        doc.add_heading("1) Interprétation des résultats", level=2)
        doc.add_heading("1.1) Multiclasse", level=3)
        doc.add_paragraph(txt_mc)
        doc.add_heading("1.2) Binaire", level=3)
        doc.add_paragraph(txt_bi)

        doc.add_heading("2) Tableaux récapitulatifs", level=2)

        def add_table(title, df):
            doc.add_heading(title, level=3)
            if df is None or len(df) == 0:
                doc.add_paragraph("Aucun résultat.")
                return

            preferred = [
                "model", "TrainTime_s",
                "Acc", "BAcc", "Macro_F1", "Weighted_F1",
                "MCC", "Kappa",
                "ROC_AUC", "PR_AUC",
                "ROC_AUC_macro", "PR_AUC_macro", "ROC_AUC_weighted", "PR_AUC_weighted",
                "TPR", "FPR", "TNR", "FNR"
            ]
            cols = [c for c in preferred if c in df.columns]
            if not cols:
                cols = list(df.columns)

            df2 = df[cols].copy()

            table = doc.add_table(rows=1, cols=len(cols))
            hdr = table.rows[0].cells
            for j, c in enumerate(cols):
                hdr[j].text = str(c)

            for _, r in df2.iterrows():
                cells = table.add_row().cells
                for j, c in enumerate(cols):
                    v = r[c]
                    if isinstance(v, float):
                        cells[j].text = f"{v:.4f}"
                    else:
                        cells[j].text = str(v)

        add_table("2.1) Multiclasse", df_mc)
        add_table("2.2) Binaire", df_bi)

        doc.add_heading("3) Annexe — Sortie console complète", level=2)
        p = doc.add_paragraph()
        run = p.add_run(console_text)
        run.font.name = "Courier New"
        run.font.size = Pt(8)

        out_path = str(base_name) + ".docx"
        doc.save(out_path)
        return out_path

    out_path = str(base_name) + ".txt"
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("UNSW-NB15 — Rapport de résultats (ELM / NN / SVM)\n")
        f.write(f"Généré le : {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
        f.write("=== Interprétation ===\n")
        f.write(txt_mc + "\n\n" + txt_bi + "\n\n")
        f.write("=== Console complète ===\n")
        f.write(console_text)
    return out_path


# =========================================================
# MAIN
# =========================================================
report_path = None
LABEL_COL = None
BENIGN_LABEL = None
time_col = None

try:
    h1("CHARGEMENT + NETTOYAGE")

    parquet_path = os.path.join(BASE_DIR, PARQUET_NAME)
    df = pd.read_parquet(parquet_path)

    df = df.replace([np.inf, -np.inf], np.nan).dropna().reset_index(drop=True)
    print(f"[INFO] Nb lignes après nettoyage: {len(df)}")

    LABEL_COL = pick_label_col(df)

    BENIGN_LABEL = BENIGN_LABEL_DEFAULT
    uniq = pd.Series(df[LABEL_COL]).dropna().unique()
    uniq = ensure_list(uniq)

    if BENIGN_LABEL not in uniq:
        if all([str(x) in ["0", "1"] for x in uniq]) or set(uniq).issubset({0, 1, np.int64(0), np.int64(1)}):
            BENIGN_LABEL = 0
        else:
            if "BENIGN" in uniq:
                BENIGN_LABEL = "BENIGN"

    time_col = detect_time_col(df)
    if time_col is not None:
        df = df.sort_values(time_col).reset_index(drop=True)

    show_counts(df[LABEL_COL], "Distribution globale (après nettoyage)", max_rows=25)

    h1("ÉCHANTILLON 'ÉTALÉ + MAJORITAIRE'")

    df_sample, top_attacks = build_majority_spread_sample(
        df, LABEL_COL, BENIGN_LABEL,
        n_total=N_SAMPLE,
        top_k_attacks=TOP_K_ATTACKS,
        benign_keep=BENIGN_KEEP,
        min_per_class=MIN_PER_CLASS,
        time_col=time_col
    )

    print(f"[INFO] Taille échantillon: {len(df_sample)}")
    show_counts(df_sample[LABEL_COL], "Distribution échantillon", max_rows=50)

    h1("MULTICLASSE — split temporel par classe (70/30)")

    df_train_mc, df_test_mc = classwise_time_split(df_sample, LABEL_COL, time_col=time_col, train_ratio=TRAIN_RATIO)

    show_counts(df_train_mc[LABEL_COL], "TRAIN multiclasse (support)", max_rows=50)
    show_counts(df_test_mc[LABEL_COL],  "TEST  multiclasse (support)", max_rows=50)

    le = LabelEncoder().fit(df_sample[LABEL_COL])
    class_names = list(le.classes_)
    y_train_mc = le.transform(df_train_mc[LABEL_COL])
    y_test_mc  = le.transform(df_test_mc[LABEL_COL])

    X_train_mc, X_test_mc = build_X(df_train_mc, df_test_mc, time_col=time_col)
    print(f"[INFO] Nb features (MC) après variance: {X_train_mc.shape[1]}")

    run_multiclass_models(X_train_mc, X_test_mc, y_train_mc, y_test_mc, class_names)

    h1("BINAIRE — split temporel par classe (70/30)")

    df_bin = df_sample.copy()
    df_bin["LabelBin"] = np.where(df_bin[LABEL_COL].values == BENIGN_LABEL, "BENIGN", "ATTACK")

    df_train_b, df_test_b = classwise_time_split(df_bin, "LabelBin", time_col=time_col, train_ratio=TRAIN_RATIO)

    show_counts(df_train_b["LabelBin"], "TRAIN binaire (support)", max_rows=10)
    show_counts(df_test_b["LabelBin"],  "TEST  binaire (support)", max_rows=10)

    y_train_bin = (df_train_b["LabelBin"].values != "BENIGN").astype(int)
    y_test_bin  = (df_test_b["LabelBin"].values  != "BENIGN").astype(int)

    X_train_b, X_test_b = build_X(df_train_b, df_test_b, time_col=time_col)
    print(f"[INFO] Nb features (BIN) après variance: {X_train_b.shape[1]}")

    run_binary_models(X_train_b, X_test_b, y_train_bin, y_test_bin)

    h1(" PIPELINE TERMINÉ")

except Exception as e:
    print("\n[ERREUR FATALE] Une exception a interrompu le pipeline.")
    print("Détails:", str(e))
    print("\nTraceback complet:")
    print(traceback.format_exc())

finally:
    # restore stdout
    sys.stdout = _original_stdout

    # Génération rapport Word/TXT (même si crash)
    try:
        meta = dict(
            dataset_path=BASE_DIR,
            parquet_name=PARQUET_NAME,
            label_col=LABEL_COL,
            benign_label=BENIGN_LABEL,
            time_col=time_col,
            n_sample=N_SAMPLE,
            train_ratio=TRAIN_RATIO,
            var_threshold=VAR_THRESHOLD
        )
        console_text = _console_buffer.getvalue()
        report_path = save_report(console_text, RESULTS, meta)

        print(f"[RAPPORT] Rapport généré : {report_path}")
        if not DOCX_OK:
            print("[RAPPORT] python-docx non installé -> rapport écrit en .txt (fallback).")
            print("Installe python-docx avec: pip install python-docx")
    except Exception as e2:
        print(f"[RAPPORT] Échec génération rapport: {e2}")
